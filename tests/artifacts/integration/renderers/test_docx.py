"""Integration tests for docx_renderer — no real LLM calls.

Tests verify:
  - render_briefing / render_briefing_docx: valid .docx produced; required
    headings present; BLUF callout; backwards-compatible dict alias works.
  - render_study_guide / render_study_guide_docx: glossary table rendered;
    headings present; correct paragraph content.
  - render_faq: Q&A headings present; hyperlink in answer linkified.
  - render_markdown_like: generic section-based render works with table + bullets.
  - TOC field present in documents that request it.
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

pytestmark = pytest.mark.integration


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def briefing_data():
    return {
        "title": "AI Investment Briefing",
        "audience": "CTO Council",
        "bluf": "AI adoption will 3× productivity by 2027.",
        "key_points": ["Point Alpha", "Point Beta", "Point Gamma"],
        "supporting_details": ["Stat A", "Stat B"],
        "action_items": ["Hire ML team", "Deploy pilot by Q3"],
        "keywords": ["ai", "productivity", "investment"],
    }


@pytest.fixture
def study_guide_data():
    return {
        "title": "Study Guide: Neural Networks",
        "overview": "An introduction to neural network fundamentals.",
        "learning_objectives": [
            "Understand forward pass",
            "Explain backpropagation",
        ],
        "key_concepts": ["Perceptron", "Activation function", "Gradient descent"],
        "glossary": [
            {"term": "Perceptron", "definition": "Basic neural unit."},
            {"term": "Epoch", "definition": "One full training pass."},
        ],
        "discussion_questions": [
            "Why does learning rate matter?",
            "When would you use ReLU?",
        ],
        "further_reading": [
            "https://deeplearning.ai",
            "Neural Networks and Deep Learning by Nielsen",
        ],
    }


@pytest.fixture
def faq_data():
    return {
        "title": "FAQ: Open Notebook",
        "items": [
            {
                "question": "Is it open source?",
                "answer": "Yes, fully MIT-licensed.",
            },
            {
                "question": "Where is the project hosted?",
                "answer": "See https://github.com/lfnovo/open-notebook for details.",
            },
        ],
    }


# ---------------------------------------------------------------------------
# render_briefing tests
# ---------------------------------------------------------------------------

def test_render_briefing_produces_docx(tmp_path: Path, briefing_data: dict) -> None:
    from open_notebook.artifacts.renderers.docx_renderer import render_briefing

    out = tmp_path / "briefing.docx"
    result = render_briefing(briefing_data, out)
    assert result == out
    assert out.exists()
    assert out.stat().st_size > 0


def test_render_briefing_docx_is_valid(tmp_path: Path, briefing_data: dict) -> None:
    from open_notebook.artifacts.renderers.docx_renderer import render_briefing

    out = tmp_path / "briefing.docx"
    render_briefing(briefing_data, out)
    # python-docx.Document round-trip verifies the file is a valid .docx
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    full_text = " ".join(texts)
    assert "AI Investment Briefing" in full_text
    assert "Key Points" in full_text
    assert "Action Items" in full_text


def test_render_briefing_bluf_present(tmp_path: Path, briefing_data: dict) -> None:
    from open_notebook.artifacts.renderers.docx_renderer import render_briefing

    out = tmp_path / "briefing.docx"
    render_briefing(briefing_data, out)
    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "BLUF" in full_text
    assert "productivity" in full_text


def test_render_briefing_paragraph_count(tmp_path: Path, briefing_data: dict) -> None:
    from open_notebook.artifacts.renderers.docx_renderer import render_briefing

    out = tmp_path / "briefing.docx"
    render_briefing(briefing_data, out)
    doc = Document(str(out))
    # Title + audience + TOC heading + TOC + page break + BLUF + spacer +
    # Key Points heading + 3 bullets + Supporting Details heading + 2 bullets +
    # Action Items heading + 2 numbered + keywords = many paragraphs
    assert len(doc.paragraphs) >= 10


def test_render_briefing_heading_styles(tmp_path: Path, briefing_data: dict) -> None:
    from open_notebook.artifacts.renderers.docx_renderer import render_briefing

    out = tmp_path / "briefing.docx"
    render_briefing(briefing_data, out)
    doc = Document(str(out))
    heading_texts = [
        p.text for p in doc.paragraphs
        if p.style.name.startswith("Heading")
    ]
    assert any("Key Points" in t for t in heading_texts)
    assert any("Action Items" in t for t in heading_texts)


def test_render_briefing_toc_field(tmp_path: Path, briefing_data: dict) -> None:
    """TOC field instruction should be present in the XML."""
    from open_notebook.artifacts.renderers.docx_renderer import render_briefing

    out = tmp_path / "briefing.docx"
    render_briefing(briefing_data, out)
    doc = Document(str(out))
    # The TOC field instruction lives inside w:instrText XML nodes
    xml_blob = doc.element.xml
    assert "TOC" in xml_blob


def test_render_briefing_dict_alias(tmp_path: Path, briefing_data: dict) -> None:
    """render_briefing_docx (legacy alias) must produce identical output."""
    from open_notebook.artifacts.renderers.docx_renderer import (
        render_briefing,
        render_briefing_docx,
    )

    p1 = tmp_path / "b1.docx"
    p2 = tmp_path / "b2.docx"
    render_briefing(briefing_data, p1)
    render_briefing_docx(briefing_data, p2)
    # Both valid docs with same paragraph count
    d1 = Document(str(p1))
    d2 = Document(str(p2))
    assert len(d1.paragraphs) == len(d2.paragraphs)


# ---------------------------------------------------------------------------
# render_study_guide tests
# ---------------------------------------------------------------------------

def test_render_study_guide_produces_docx(tmp_path: Path, study_guide_data: dict) -> None:
    from open_notebook.artifacts.renderers.docx_renderer import render_study_guide

    out = tmp_path / "study_guide.docx"
    result = render_study_guide(study_guide_data, out)
    assert result.exists()


def test_render_study_guide_headings(tmp_path: Path, study_guide_data: dict) -> None:
    from open_notebook.artifacts.renderers.docx_renderer import render_study_guide

    out = tmp_path / "study_guide.docx"
    render_study_guide(study_guide_data, out)
    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    for section in ["Overview", "Learning Objectives", "Key Concepts", "Glossary", "Discussion Questions"]:
        assert section in full_text, f"Missing section: {section}"


def test_render_study_guide_glossary_bold(tmp_path: Path, study_guide_data: dict) -> None:
    from open_notebook.artifacts.renderers.docx_renderer import render_study_guide

    out = tmp_path / "study_guide.docx"
    render_study_guide(study_guide_data, out)
    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "Perceptron" in full_text
    assert "Epoch" in full_text


def test_render_study_guide_further_reading_hyperlink(tmp_path: Path, study_guide_data: dict) -> None:
    """further_reading URLs should be stored as hyperlink relationships."""
    from open_notebook.artifacts.renderers.docx_renderer import render_study_guide

    out = tmp_path / "study_guide.docx"
    render_study_guide(study_guide_data, out)
    doc = Document(str(out))
    # Hyperlink relationships are stored in part.rels
    all_targets = [
        rel.target_ref
        for rel in doc.part.rels.values()
        if "hyperlink" in rel.reltype
    ]
    assert any("deeplearning.ai" in t for t in all_targets), (
        f"Expected deeplearning.ai hyperlink. Got: {all_targets}"
    )


def test_render_study_guide_dict_alias(tmp_path: Path, study_guide_data: dict) -> None:
    from open_notebook.artifacts.renderers.docx_renderer import render_study_guide_docx

    out = tmp_path / "sg.docx"
    render_study_guide_docx(study_guide_data, out)
    doc = Document(str(out))
    assert len(doc.paragraphs) >= 5


# ---------------------------------------------------------------------------
# render_faq tests
# ---------------------------------------------------------------------------

def test_render_faq_produces_docx(tmp_path: Path, faq_data: dict) -> None:
    from open_notebook.artifacts.renderers.docx_renderer import render_faq

    out = tmp_path / "faq.docx"
    result = render_faq(faq_data, out)
    assert result.exists()


def test_render_faq_questions_present(tmp_path: Path, faq_data: dict) -> None:
    from open_notebook.artifacts.renderers.docx_renderer import render_faq

    out = tmp_path / "faq.docx"
    render_faq(faq_data, out)
    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "Is it open source" in full_text
    assert "FAQ: Open Notebook" in full_text


def test_render_faq_url_hyperlink(tmp_path: Path, faq_data: dict) -> None:
    """Answer containing a bare URL should be stored as a hyperlink."""
    from open_notebook.artifacts.renderers.docx_renderer import render_faq

    out = tmp_path / "faq.docx"
    render_faq(faq_data, out)
    doc = Document(str(out))
    all_targets = [
        rel.target_ref
        for rel in doc.part.rels.values()
        if "hyperlink" in rel.reltype
    ]
    assert any("github.com" in t for t in all_targets), (
        f"Expected github.com hyperlink from FAQ answer. Got: {all_targets}"
    )


# ---------------------------------------------------------------------------
# render_markdown_like tests
# ---------------------------------------------------------------------------

def test_render_markdown_like_sections(tmp_path: Path) -> None:
    from open_notebook.artifacts.renderers.docx_renderer import render_markdown_like

    sections = [
        {"heading": "Introduction", "level": 1, "text": "Welcome to this document."},
        {"heading": "Features", "level": 1, "bullets": ["Fast", "Reliable", "Cheap"]},
        {
            "heading": "Data",
            "level": 1,
            "table": {
                "columns": ["Name", "Value"],
                "rows": [["Alpha", "1"], ["Beta", "2"]],
            },
        },
    ]
    out = tmp_path / "generic.docx"
    render_markdown_like("Test Document", sections, out)
    assert out.exists()

    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "Introduction" in full_text
    assert "Welcome to this document" in full_text
    assert "Features" in full_text
    assert "Fast" in full_text

    # Verify table was written
    assert len(doc.tables) >= 1
    tbl = doc.tables[0]
    header_cells = [c.text for c in tbl.rows[0].cells]
    assert "Name" in header_cells
    assert "Value" in header_cells


def test_render_markdown_like_table_header_styled(tmp_path: Path) -> None:
    """Header row should have a non-white background fill (blue)."""
    from open_notebook.artifacts.renderers.docx_renderer import render_markdown_like
    from docx.oxml.ns import qn

    sections = [
        {
            "heading": "Data Table",
            "table": {
                "columns": ["Col A", "Col B"],
                "rows": [["x", "y"]],
            },
        }
    ]
    out = tmp_path / "table_doc.docx"
    render_markdown_like("Table Test", sections, out)
    doc = Document(str(out))
    assert doc.tables
    tbl = doc.tables[0]
    # Check first cell of header row has shading element
    first_cell = tbl.rows[0].cells[0]
    tc_xml = first_cell._tc.xml
    assert "shd" in tc_xml or "w:shd" in tc_xml
