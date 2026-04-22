"""python-docx renderers for prose-style artifacts.

Public API
----------
render_briefing(schema, output_path)    -> Path
render_study_guide(schema, output_path) -> Path
render_faq(schema, output_path)         -> Path
render_markdown_like(title, sections, output_path) -> Path

Legacy dict-based aliases kept for backwards compatibility:
render_briefing_docx(data, path)    -> Path
render_study_guide_docx(data, path) -> Path
"""
from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


# ---------------------------------------------------------------------------
# Typography constants
# ---------------------------------------------------------------------------

_BODY_FONT = "Calibri"
_HEADING_FONT = "Calibri Light"
_BODY_PT = 11
_H1_PT = 18
_H2_PT = 14
_H3_PT = 12
_LINE_SPACING = 1.15

# Brand colour for heading underlines / table headers
_HEADING_COLOR = RGBColor(0x1F, 0x49, 0x7D)   # Word "Heading 1" blue
_ACCENT_COLOR = RGBColor(0x2E, 0x74, 0xB5)    # Lighter blue
_TABLE_HEADER_BG = "2E74B5"                   # hex string for XML


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _apply_body_style(doc: Document) -> None:
    """Set Normal style to Calibri 11 pt with 1.15 line spacing."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = _BODY_FONT
    font.size = Pt(_BODY_PT)
    pf = style.paragraph_format
    pf.line_spacing = _LINE_SPACING


def _styled_heading(doc: Document, text: str, level: int) -> None:
    """Add a heading with Calibri Light at the specified pt size."""
    if not text:
        return
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = _HEADING_FONT
        run.font.color.rgb = _HEADING_COLOR
        run.font.size = Pt({0: 22, 1: _H1_PT, 2: _H2_PT, 3: _H3_PT}.get(level, _H3_PT))


def _add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        if item is None:
            continue
        doc.add_paragraph(str(item), style="List Bullet")


def _add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        if item is None:
            continue
        doc.add_paragraph(str(item), style="List Number")


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Insert a clickable hyperlink into *paragraph*."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _insert_toc(doc: Document) -> None:
    """Insert a Word TOC field (rendered by Word on open)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Right-click → Update Field to refresh this Table of Contents."
    placeholder.append(t)
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(placeholder)
    run._r.append(fld_char_end)
    doc.add_paragraph()  # spacer after TOC


def _set_table_header_style(table) -> None:
    """Bold + blue background on the first row of a table."""
    hdr_row = table.rows[0]
    for cell in hdr_row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), _TABLE_HEADER_BG)
        tcPr.append(shd)


def _add_callout_box(doc: Document, label: str, content: str) -> None:
    """Mimics a callout with a bold label + indented paragraph."""
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = _ACCENT_COLOR
    p.add_run(content)
    p.paragraph_format.left_indent = Pt(18)


def _save(doc: Document, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Generic section-based renderer
# ---------------------------------------------------------------------------

def render_markdown_like(
    title: str,
    sections: List[Dict[str, Any]],
    output_path: Path,
    *,
    include_toc: bool = True,
) -> Path:
    """Render a document from a title + list of section dicts.

    Each section dict supports:
      heading (str)  — section title
      level (int)    — 1 / 2 / 3 (default 1)
      text (str)     — prose paragraph
      bullets (list) — bullet list items
      numbered (list)— numbered list items
      table (dict)   — {"columns": [...], "rows": [[...]...]}
    """
    doc = Document()
    _apply_body_style(doc)
    _styled_heading(doc, title, level=0)

    if include_toc:
        doc.add_paragraph("Table of Contents", style="Heading 2")
        _insert_toc(doc)
        doc.add_page_break()

    for section in sections:
        heading = section.get("heading")
        level = int(section.get("level", 1))
        if heading:
            _styled_heading(doc, heading, level=level)
        if text := section.get("text"):
            doc.add_paragraph(str(text))
        if bullets := section.get("bullets"):
            _add_bullets(doc, bullets)
        if numbered := section.get("numbered"):
            _add_numbered(doc, numbered)
        if tbl := section.get("table"):
            cols = tbl.get("columns", [])
            rows = tbl.get("rows", [])
            if cols and rows:
                t = doc.add_table(rows=1 + len(rows), cols=len(cols))
                t.style = "Table Grid"
                for ci, col in enumerate(cols):
                    t.rows[0].cells[ci].text = str(col)
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row[:len(cols)]):
                        t.rows[ri + 1].cells[ci].text = str(val)
                _set_table_header_style(t)

    return _save(doc, output_path)


# ---------------------------------------------------------------------------
# Briefing renderer
# ---------------------------------------------------------------------------

def render_briefing(schema, output_path: Path) -> Path:
    """Render a BriefingSchema (or dict) to a .docx file.

    Accepts either a Pydantic BriefingSchema instance or a plain dict.
    """
    data: Dict[str, Any] = schema if isinstance(schema, dict) else schema.model_dump()

    doc = Document()
    _apply_body_style(doc)

    _styled_heading(doc, data.get("title", "Executive Briefing"), level=0)

    # Audience / metadata row
    audience = data.get("audience")
    if audience:
        p = doc.add_paragraph()
        r = p.add_run("Audience: ")
        r.bold = True
        r.font.color.rgb = _ACCENT_COLOR
        p.add_run(str(audience)).italic = True

    doc.add_paragraph("Table of Contents", style="Heading 2")
    _insert_toc(doc)
    doc.add_page_break()

    # BLUF callout
    bluf = data.get("bluf")
    if bluf:
        _add_callout_box(doc, "BLUF", str(bluf))
        doc.add_paragraph()  # spacer

    if kp := data.get("key_points"):
        _styled_heading(doc, "Key Points", level=1)
        _add_bullets(doc, kp)

    if sd := data.get("supporting_details"):
        _styled_heading(doc, "Supporting Details", level=1)
        _add_bullets(doc, sd)

    if ai := data.get("action_items"):
        _styled_heading(doc, "Action Items", level=1)
        _add_numbered(doc, ai)

    if keywords := data.get("keywords"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("Keywords: ")
        r.bold = True
        p.add_run(", ".join(str(k) for k in keywords)).italic = True

    return _save(doc, output_path)


# ---------------------------------------------------------------------------
# Study guide renderer
# ---------------------------------------------------------------------------

def render_study_guide(schema, output_path: Path) -> Path:
    """Render a StudyGuideSchema (or dict) to a .docx file."""
    data: Dict[str, Any] = schema if isinstance(schema, dict) else schema.model_dump()

    doc = Document()
    _apply_body_style(doc)

    title = data.get("title", "Study Guide")
    _styled_heading(doc, title, level=0)

    doc.add_paragraph("Table of Contents", style="Heading 2")
    _insert_toc(doc)
    doc.add_page_break()

    if overview := data.get("overview"):
        _styled_heading(doc, "Overview", level=1)
        doc.add_paragraph(str(overview))

    if lo := data.get("learning_objectives"):
        _styled_heading(doc, "Learning Objectives", level=1)
        _add_numbered(doc, lo)

    if kc := data.get("key_concepts"):
        _styled_heading(doc, "Key Concepts", level=1)
        _add_bullets(doc, kc)

    glossary = data.get("glossary") or []
    if glossary:
        _styled_heading(doc, "Glossary", level=1)
        for item in glossary:
            if isinstance(item, dict):
                term = item.get("term", "")
                defn = item.get("definition", "")
            else:
                term = str(item)
                defn = ""
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(str(term))
            r.bold = True
            p.add_run(f" — {defn}")

    if dq := data.get("discussion_questions"):
        _styled_heading(doc, "Discussion Questions", level=1)
        _add_numbered(doc, dq)

    # Further reading with hyperlinks where values look like URLs
    fr = data.get("further_reading") or []
    if fr:
        _styled_heading(doc, "Further Reading", level=1)
        for item in fr:
            text = str(item)
            p = doc.add_paragraph(style="List Bullet")
            if text.startswith(("http://", "https://")):
                _add_hyperlink(p, text, text)
            else:
                p.add_run(text)

    return _save(doc, output_path)


# ---------------------------------------------------------------------------
# FAQ renderer
# ---------------------------------------------------------------------------

def render_faq(schema, output_path: Path) -> Path:
    """Render a FAQSchema (or dict) to a .docx file."""
    data: Dict[str, Any] = schema if isinstance(schema, dict) else schema.model_dump()

    doc = Document()
    _apply_body_style(doc)

    _styled_heading(doc, data.get("title", "FAQ"), level=0)

    doc.add_paragraph("Table of Contents", style="Heading 2")
    _insert_toc(doc)
    doc.add_page_break()

    for item in data.get("items", []):
        if isinstance(item, dict):
            question = item.get("question", "")
            answer = item.get("answer", "")
        else:
            question = str(item)
            answer = ""
        _styled_heading(doc, str(question), level=2)
        if answer:
            # Detect URLs in the answer and linkify trailing bare URLs
            words = str(answer).split()
            if any(w.startswith(("http://", "https://")) for w in words):
                p = doc.add_paragraph()
                for word in words:
                    if word.startswith(("http://", "https://")):
                        _add_hyperlink(p, word + " ", word)
                    else:
                        p.add_run(word + " ")
            else:
                doc.add_paragraph(str(answer))

    return _save(doc, output_path)


# ---------------------------------------------------------------------------
# Backwards-compatible aliases (used by existing generators)
# ---------------------------------------------------------------------------

def render_briefing_docx(data: Dict[str, Any], path: Path) -> Path:  # noqa: D401
    """Legacy dict-based alias — delegates to render_briefing."""
    return render_briefing(data, path)


def render_study_guide_docx(data: Dict[str, Any], path: Path) -> Path:  # noqa: D401
    """Legacy dict-based alias — delegates to render_study_guide."""
    return render_study_guide(data, path)
